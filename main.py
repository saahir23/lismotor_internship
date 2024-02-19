from selenium import webdriver
from bs4 import BeautifulSoup
import requests
import json
from docx import Document

# Initialize Selenium WebDriver
options = webdriver.ChromeOptions()
options.add_argument('--headless')  
driver = webdriver.Chrome(options=options)

def fetch_content(url):
    needs_javascript = False  
    if needs_javascript:
        driver.get(url)
        html = driver.page_source
    else:
        response = requests.get(url)
        html = response.content
    return html

def scrape_url(url):
    html = fetch_content(url)
    soup = BeautifulSoup(html, 'html.parser')
    texts = soup.find('body').find_all(text=True)
    content = ' '.join(text.strip() for text in texts if text.parent.name not in ['script', 'style', 'head', 'title', 'meta', '[document]'])
    return {
        'url': url,
        'content': content  
    }

def load_urls(filename='search_results.json'):
    with open(filename, 'r') as file:
        data = json.load(file)
    return data

def save_to_word(data, filename='scraped_results.docx'):
    doc = Document()
    for entry in data:
        query = entry['query']
        doc.add_paragraph(f"Query: {query}")
        
        for url in entry['urls']:
            try:
                result = scrape_url(url)
                doc.add_paragraph(f"URL: {url}")

                doc.add_paragraph(f"Content: {result['content'][:500]}...")  
            except Exception as e:
                print(f"Error scraping {url}: {e}")
                a
            doc.add_paragraph("") 

    doc.save(filename)

if __name__ == "__main__":
    data = load_urls()  
    save_to_word(data)  

    driver.quit()
